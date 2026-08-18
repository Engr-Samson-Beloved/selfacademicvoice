"""Table content must be rewritten like any other prose.

rewrite_docx() iterated doc.paragraphs, which in python-docx returns only
body-level paragraphs. Everything inside a <w:tbl> was invisible to it, so
table prose passed through byte-identical however it read.

llm.ask is stubbed: no API keys, no quota.
"""

import io
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import docx  # noqa: E402

from ghostwriter import llm, rewrite  # noqa: E402

CELL_PROSE = (
    "The proposed system captures live video from a standard webcam and then "
    "classifies each detected hand pose into a matching playback command."
)
BODY_PROSE = (
    "Public libraries are established to provide timely information resources "
    "to every member of the surrounding community without any charge."
)
# Distinct from CELL_PROSE on purpose: the pipeline drops near-duplicate
# paragraphs, so reusing the same text would hide whether nesting was traversed.
NESTED_PROSE = (
    "Landmark coordinates are extracted for each frame and compared against the "
    "stored gesture templates before any playback command is finally dispatched."
)


def _build_doc(nested=False, merged=False):
    d = docx.Document()
    d.add_paragraph(BODY_PROSE)
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Stage"
    t.cell(0, 1).text = "Description"
    t.cell(1, 0).text = "Capture"
    t.cell(1, 1).text = CELL_PROSE
    if merged:
        t2 = d.add_table(rows=1, cols=2)
        t2.cell(0, 0).merge(t2.cell(0, 1)).text = CELL_PROSE
    if nested:
        outer = d.add_table(rows=1, cols=1)
        inner = outer.cell(0, 0).add_table(rows=1, cols=1)
        inner.cell(0, 0).text = NESTED_PROSE
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _all_text(blob):
    d = docx.Document(io.BytesIO(blob))
    parts = [p.text for p in d.paragraphs]

    def walk(tables):
        for t in tables:
            for row in t.rows:
                for c in row.cells:
                    parts.extend(p.text for p in c.paragraphs)
                    walk(c.tables)

    walk(d.tables)
    return "\n".join(parts)


class StubLLM:
    def __init__(self):
        self.seen = []

    def __call__(self, prompt, system_prompt=None, temperature=None):
        self.seen.append(prompt)
        nums = [int(m) for m in re.findall(r"^(\d+)\.", prompt, re.MULTILINE)]
        return "\n".join(
            f"{n}. Rewritten line number {n} carrying the same meaning across." for n in nums
        )


def _run(blob):
    stub = StubLLM()
    original = llm.ask
    llm.ask = stub
    try:
        return rewrite.rewrite_docx(blob, "system"), stub
    finally:
        llm.ask = original


def test_iter_paragraphs_sees_table_cells():
    d = docx.Document(io.BytesIO(_build_doc()))
    body_only = {id(p._p) for p in d.paragraphs}
    everything = {id(p._p) for p in rewrite.iter_paragraphs(d)}
    assert len(everything) > len(body_only), "table paragraphs still not visible"
    assert body_only <= everything, "body paragraphs went missing"


def test_table_prose_is_sent_to_the_model():
    _, stub = _run(_build_doc())
    joined = "\n".join(stub.seen)
    assert "standard webcam" in joined, "cell prose was never sent for rewriting"


def test_table_prose_is_actually_changed():
    out, _ = _run(_build_doc())
    assert CELL_PROSE not in _all_text(out), "cell prose survived byte-identical"


def test_nested_table_prose_is_rewritten():
    out, stub = _run(_build_doc(nested=True))
    assert "stored gesture templates" in "\n".join(stub.seen), \
        "nested cell prose was never sent for rewriting"
    assert NESTED_PROSE not in _all_text(out), "nested cell prose survived byte-identical"


def test_merged_cell_not_rewritten_twice():
    """A merged cell appears once per spanned column in row.cells."""
    d = docx.Document(io.BytesIO(_build_doc(merged=True)))
    seen = [id(p._p) for p in rewrite.iter_paragraphs(d)]
    assert len(seen) == len(set(seen)), "a paragraph was yielded more than once"


def test_document_still_valid_and_table_intact():
    out, _ = _run(_build_doc())
    d = docx.Document(io.BytesIO(out))
    assert len(d.tables) == 1, "table was lost"
    assert len(d.tables[0].rows) == 2 and len(d.tables[0].columns) == 2
    assert d.tables[0].cell(0, 0).text.strip(), "header cell emptied"


def test_short_header_cells_left_alone():
    """One-word headers are not prose and must not be paraphrased."""
    out, _ = _run(_build_doc())
    text = _all_text(out)
    assert "Stage" in text and "Description" in text, "short headers were altered"


def test_terse_data_cells_are_not_paraphrased():
    """Regression: the first table pass treated every cell as prose and produced
    "Ultra-High Frequency which is called UHF." and "Tracking of Vehicle and
    long range." from what were column labels and matrix entries."""
    d = docx.Document()
    d.add_paragraph(BODY_PROSE)
    t = d.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Band"
    t.cell(0, 1).text = "Energy Savings Potential"
    t.cell(1, 0).text = "Ultra-High Frequency (UHF)"
    t.cell(1, 1).text = "Access control, animal tagging"
    t.cell(2, 0).text = "High Frequency (HF)"
    t.cell(2, 1).text = "Vehicle & long-range tracking"
    buf = io.BytesIO()
    d.save(buf)

    out, _ = _run(buf.getvalue())
    text = _all_text(out)
    for original in ("Ultra-High Frequency (UHF)", "Energy Savings Potential",
                     "Access control, animal tagging", "Vehicle & long-range tracking",
                     "High Frequency (HF)"):
        assert original in text, f"terse cell was paraphrased: {original!r}"


def test_header_row_never_rewritten_even_when_wordy():
    d = docx.Document()
    d.add_paragraph(BODY_PROSE)
    t = d.add_table(rows=2, cols=1)
    header = ("Comparative assessment of the available automation approaches "
              "across installation cost and achievable savings")
    t.cell(0, 0).text = header
    t.cell(1, 0).text = CELL_PROSE
    buf = io.BytesIO()
    d.save(buf)

    out, _ = _run(buf.getvalue())
    text = _all_text(out)
    assert header in text, "header row was rewritten"
    assert CELL_PROSE not in text, "body cell prose should still be rewritten"


if __name__ == "__main__":
    tests = [(n, o) for n, o in sorted(globals().items())
             if n.startswith("test_") and callable(o)]
    failed = 0
    for name, fn in tests:
        try:
            fn()
            print(f"PASS  {name}")
        except AssertionError as e:
            failed += 1
            print(f"FAIL  {name}: {e}")
        except Exception as e:
            failed += 1
            print(f"ERROR {name}: {type(e).__name__}: {e}")
    print(f"\n{len(tests) - failed}/{len(tests)} passed")
    raise SystemExit(1 if failed else 0)
