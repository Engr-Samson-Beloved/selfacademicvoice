"""Tests for the measured voice profile and the rewrite-time voice gate.

Runs without API keys: ``llm.ask`` is replaced with a stub, so no Gemini or Groq
call is made and no quota is consumed.

    python tests/test_voice_gate.py     # or: pytest tests/test_voice_gate.py
"""

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from ghostwriter import config, llm, rewrite, voiceprofile  # noqa: E402

PROFILE = voiceprofile.load_json(Path("data/voice_profile.json"))


# --------------------------------------------------------------- unit: rules

def test_banned_connective_flagged():
    orig = "Public libraries help the community."
    bad = "However, public libraries assist the community in many ways."
    problems = voiceprofile.sentence_violations(PROFILE, bad, orig)
    assert any("However" in p for p in problems), problems


def test_attested_connector_not_flagged():
    # "Furthermore" IS in the corpus, so it must survive the gate.
    orig = "Public libraries help the community."
    ok = "Furthermore, public libraries assist the community."
    problems = voiceprofile.sentence_violations(PROFILE, ok, orig)
    assert not any("Furthermore" in p for p in problems), problems


def test_semicolon_and_dash_flagged():
    orig = "The library opened in 1956."
    assert voiceprofile.sentence_violations(
        PROFILE, "The library opened in 1956; it moved later.", orig)
    assert voiceprofile.sentence_violations(
        PROFILE, "The library opened in 1956 — it moved later.", orig)


def test_length_inflation_flagged():
    orig = "Education depends on the right information."
    inflated = " ".join(["word"] * 80)
    problems = voiceprofile.sentence_violations(PROFILE, inflated, orig)
    assert any("shorten" in p for p in problems), problems


def test_comma_inflation_flagged():
    orig = "The library serves students."
    bad = ("The library, which is large, serves students, teachers, scholars, "
           "officials, and dropouts, among others.")
    problems = voiceprofile.sentence_violations(PROFILE, bad, orig)
    assert any("commas" in p for p in problems), problems


def test_content_deletion_flagged():
    """Regression: only length INFLATION was checked, so the model compressed
    instead of rewriting - a 68-word sentence came back as 11, silently."""
    orig = ("A Gesture-Controlled Music Player is one practical realisation of this "
            "idea, because instead of pressing physical buttons, tapping a touchscreen "
            "or clicking a mouse, the user simply raises a hand in front of the camera "
            "and the system reads the pose and dispatches the matching command.")
    compressed = "One practical application of this concept is a Gesture-Controlled Music Player."
    problems = voiceprofile.sentence_violations(PROFILE, compressed, orig)
    assert any("dropped content" in p for p in problems), problems


def test_short_original_not_flagged_for_shrinkage():
    """Short sentences legitimately get shorter; the floor only applies to
    originals long enough for the loss to mean something."""
    orig = "Education depends on information."
    problems = voiceprofile.sentence_violations(PROFILE, "Learning needs information.", orig)
    assert not any("dropped content" in p for p in problems), problems


def test_verbatim_run_flagged():
    """A long copied run reads as 'not rewritten' even at moderate token
    overlap. One real sentence carried 25 consecutive words through at 0.85."""
    orig = ("The specific problem addressed in this seminar can be stated as follows: "
            "design and evaluate a system that captures live video from a standard "
            "webcam and classifies the resulting hand poses.")
    copied = ("The specific problem addressed in this seminar can be stated as follows: "
              "design and evaluate a System that reads webcam video and sorts hand poses.")
    problems = voiceprofile.sentence_violations(PROFILE, copied, orig)
    assert any("word for word" in p for p in problems), problems


def test_genuine_rewrite_has_no_long_verbatim_run():
    orig = "Public libraries are established to provide timely information resources."
    good = "Prompt information materials reach the community through Public Library."
    assert not any("word for word" in p
                   for p in voiceprofile.sentence_violations(PROFILE, good, orig))


def test_label_colon_bullet_flagged():
    """The corpus contains no "Label: explanation" bullets in prose, but the
    rewrite prompt used to mandate keeping such labels verbatim - leaving them
    simultaneously unrewritten and off-voice."""
    for s in ("Touchless operation: removes the need to touch common surfaces.",
              "• Sensitivity to lighting conditions: accuracy degrades in poor light.",
              "Extensibility: the same pipeline can be repurposed for other applications."):
        problems = voiceprofile.sentence_violations(PROFILE, s, s + " padding words here")
        assert any("Label: explanation" in p for p in problems), s


def test_ordinary_colon_lead_in_not_flagged():
    """A clause ending in a colon is not a label; only short noun phrases are."""
    for s in ("The specific problem can be stated as follows: design and evaluate a system.",
              "The study shows that majority of users visit public libraries regularly."):
        problems = voiceprofile.sentence_violations(PROFILE, s, s + " padding words here")
        assert not any("Label: explanation" in p for p in problems), s


def test_label_prefix_extraction():
    assert voiceprofile.label_prefix("Touchless operation: removes the need") == "Touchless operation"
    assert voiceprofile.label_prefix("• Extensibility: the pipeline adapts") == "Extensibility"
    assert voiceprofile.label_prefix("The specific problem can be stated as follows: design it") is None
    assert voiceprofile.label_prefix("Public libraries provide information resources.") is None


def test_clean_sentence_passes():
    orig = "Education solely depends on the availability of information."
    good = "Education depends only on having the right information at the right time."
    assert voiceprofile.sentence_violations(PROFILE, good, orig) == []


def test_profile_json_roundtrip(tmp_path=Path("data")):
    reloaded = voiceprofile.load_json(tmp_path / "voice_profile.json")
    assert reloaded.median_len == PROFILE.median_len
    assert reloaded.connectors and isinstance(reloaded.connectors[0], tuple)


# ------------------------------------------------------ regressions: the audit

def _comma_row(text):
    return next(r for r in voiceprofile.compare(PROFILE, text)
                if r[0] == "commas per sentence")


def test_check_flags_comma_undershoot():
    """Regression: the comma check was one-sided, so output using 0.18 commas
    per sentence against an author who uses 1.56 passed silently."""
    clipped = " ".join(
        "The system detects the hand and maps it to a control instruction." for _ in range(12)
    )
    label, author_val, cand_val, ok = _comma_row(clipped)
    assert not ok, f"undershoot not flagged (author {author_val}, candidate {cand_val})"


def test_check_still_flags_comma_overshoot():
    overloaded = " ".join(
        "The system, which is large, detects the hand, maps it, and, in turn, acts."
        for _ in range(12)
    )
    _, _, _, ok = _comma_row(overloaded)
    assert not ok, "overshoot no longer flagged"


def test_check_passes_matching_comma_rate():
    """The author's own corpus must not trip the two-sided check."""
    _, _, _, ok = _comma_row(voiceprofile.read_corpus(Path("myvoice")))
    assert ok, "author corpus fails its own comma check"


def test_read_corpus_reads_docx():
    """Regression: .docx fell through to read_text(), so --check on the
    pipeline's own output measured raw zip bytes instead of the document."""
    import docx

    tmp = Path("data/_test_reader.docx")
    d = docx.Document()
    d.add_paragraph("Public libraries are established to provide information resources.")
    d.add_paragraph("Education depends on the right information at the right time.")
    d.save(tmp)
    try:
        text = voiceprofile.read_corpus(tmp)
        assert "Public libraries" in text, f"docx not parsed, got: {text[:60]!r}"
        assert "PK" not in text[:4], "read raw zip bytes instead of document text"
    finally:
        tmp.unlink(missing_ok=True)


def test_read_corpus_rejects_unknown_suffix():
    tmp = Path("data/_test_reader.bin")
    tmp.write_bytes(b"\x00\x01\x02")
    try:
        try:
            voiceprofile.read_corpus(tmp)
        except SystemExit:
            return
        raise AssertionError("unknown suffix was not rejected")
    finally:
        tmp.unlink(missing_ok=True)


# --------------------------------------------------------- references handling

def test_references_heading_detected():
    for t in ("REFERENCES", "References", "Bibliography", "WORKS CITED",
              "7. REFERENCES", "References:", "Reference List", "### References"):
        assert rewrite.is_references_heading(t), t


def test_references_heading_not_overmatched():
    for t in ("The references show that gesture control works well",
              "METHODOLOGY", "1.2 Problem Definition",
              "This section references prior work on hand tracking"):
        assert not rewrite.is_references_heading(t), t


def test_reference_list_is_not_rewritten():
    """Reference entries must survive byte-identical.

    They were previously fed to the rewriter, which either returned them
    unchanged (journal names are not rewordable) or corrupted them - retitling
    published papers and injecting the author's coinages into citations.
    """
    citation = ("Zhang, F., Bazarevsky, V., & Grundmann, M. (2020). MediaPipe Hands: "
                "On-device real-time hand tracking (arXiv:2006.10214). arXiv.")
    document = (
        "Public libraries are established to provide timely information resources "
        "to all members of the community.\n\n"
        "REFERENCES\n\n" + citation
    )
    replies = {1: "Public libraries exist so that prompt resources reach the community."}

    class LLM(FakeLLM):
        def __call__(self, prompt, system_prompt=None, temperature=None):
            assert "MediaPipe Hands" not in prompt, "citation was sent to the model"
            return super().__call__(prompt, system_prompt, temperature)

    fake = LLM(replies)
    original_ask = llm.ask
    llm.ask = fake
    try:
        out = rewrite.rewrite_document(document, "system")
    finally:
        llm.ask = original_ask

    assert citation in out, "reference entry was altered or dropped"
    assert "REFERENCES" in out, "reference heading was dropped"


# ------------------------------------------------------- integration: the gate

class FakeLLM:
    """Stands in for llm.ask. Replies to whatever numbers the prompt asks for."""

    def __init__(self, replies):
        self.replies = replies
        self.calls = 0

    def __call__(self, prompt, system_prompt=None, temperature=None):
        self.calls += 1
        wanted = [int(m) for m in re.findall(r"^(\d+)\.", prompt, re.MULTILINE)]
        return "\n".join(
            f"{n}. {self.replies[n]}" for n in wanted if n in self.replies
        )


def _run_gate(rewrites, original_map, replies):
    fake = FakeLLM(replies)
    original_ask = llm.ask
    llm.ask = fake
    try:
        rewrite._apply_voice_gate(rewrites, original_map, "system")
    finally:
        llm.ask = original_ask
    return fake


def test_gate_repairs_a_breach():
    original_map = {1: "Public libraries help the community."}
    rewrites = {1: "However, public libraries assist the community greatly."}
    fake = _run_gate(rewrites, original_map,
                     {1: "Public libraries give real assistance to the community."})
    assert fake.calls == 1
    assert "However" not in rewrites[1]


def test_gate_rejects_a_non_improvement():
    """A 'fix' that still breaches must be discarded, not written back."""
    original_map = {1: "Public libraries help the community."}
    breaching = "However, public libraries assist the community greatly."
    rewrites = {1: breaching}
    _run_gate(rewrites, original_map,
              {1: "Moreover, public libraries assist the community greatly."})
    assert rewrites[1] == breaching, "a still-breaching repair was accepted"


def test_gate_rejects_repair_that_copies_the_original():
    """The voice gate must never undo the similarity gate."""
    orig = "Public libraries are established to provide timely information resources."
    rewrites = {1: "However, public libraries exist to supply prompt information."}
    _run_gate(rewrites, original_map={1: orig}, replies={1: orig})
    assert rewrites[1] != orig, "gate reintroduced a near-copy of the source"


def test_gate_is_noop_without_a_profile():
    """No profile file -> pipeline behaves exactly as before the gate existed."""
    saved = (rewrite._profile_cache, rewrite._profile_loaded, config.VOICE_PROFILE_FILE)
    rewrite._profile_cache, rewrite._profile_loaded = None, False
    config.VOICE_PROFILE_FILE = Path("data/does_not_exist.json")
    try:
        rewrites = {1: "However, this clearly breaches the voice rules."}
        fake = _run_gate(rewrites, {1: "Short original."}, {1: "irrelevant"})
        assert fake.calls == 0
        assert rewrites[1] == "However, this clearly breaches the voice rules."
    finally:
        rewrite._profile_cache, rewrite._profile_loaded, config.VOICE_PROFILE_FILE = saved


def test_gate_converges_and_stops():
    """When no repair is accepted the gate must stop, not spend every attempt."""
    original_map = {1: "Public libraries help the community."}
    rewrites = {1: "However, public libraries assist the community greatly."}
    fake = _run_gate(rewrites, original_map,
                     {1: "Moreover, public libraries assist the community greatly."})
    assert fake.calls == 1, f"expected one pass then stop, made {fake.calls}"


def test_full_rewrite_document_runs_with_gate():
    document = (
        "Public libraries are established to provide timely information resources "
        "to every member of the community. Education depends on the right "
        "information at the right time. The study shows that most users visit the "
        "library regularly."
    )
    replies = {
        1: "Public libraries exist so that prompt information resources reach all community members.",
        2: "Learning rests upon having correct information exactly when it is needed.",
        3: "Findings show most readers come to the library on a regular basis.",
    }

    class AnyNumberLLM(FakeLLM):
        def __call__(self, prompt, system_prompt=None, temperature=None):
            self.calls += 1
            wanted = [int(m) for m in re.findall(r"^(\d+)\.", prompt, re.MULTILINE)]
            return "\n".join(
                f"{n}. {self.replies.get(n, 'A short rewritten line for this item.')}"
                for n in wanted
            )

    fake = AnyNumberLLM(replies)
    original_ask = llm.ask
    llm.ask = fake
    try:
        out = rewrite.rewrite_document(document, "system")
    finally:
        llm.ask = original_ask

    assert out.strip(), "rewrite produced no output"
    assert fake.calls >= 1
    assert ";" not in out


if __name__ == "__main__":
    tests = [(n, o) for n, o in sorted(globals().items())
             if n.startswith("test_") and callable(o)]
    failed = 0
    for name, fn in tests:
        try:
            fn()
            print(f"PASS  {name}")
        except AssertionError as e:
            failed += 1
            print(f"FAIL  {name}: {e}")
        except Exception as e:
            failed += 1
            print(f"ERROR {name}: {type(e).__name__}: {e}")
    print(f"\n{len(tests) - failed}/{len(tests)} passed")
    raise SystemExit(1 if failed else 0)
