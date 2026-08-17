import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_CAPTION_RE = re.compile(r"^(?:Table|Figure|Fig\.?|Illustration)\b", re.IGNORECASE)
_INLINE_CAPTION_RE = re.compile(r"\b(?:Figure|Table|Fig\.?|Illustration)\s+\d+", re.IGNORECASE)
_HEADING_TEXT_RE = re.compile(r"^(?:CHAPTER\b|\d+(?:\.\d+)*\s+[A-Z])")
_MAJOR_HEADING_RE = re.compile(r"^CHAPTER\b", re.IGNORECASE)
_REFERENCE_HEADING_RE = re.compile(r"^(?:(?:LIST OF )?REFERENCES|BIBLIOGRAPHY|WORKS CITED)\b", re.IGNORECASE)
_TOC_RE = re.compile(r"^TABLE\s+OF\s+CONTENTS\b", re.IGNORECASE)
_LISTING_HEADING_RE = re.compile(
    r"^(?:(?:LIST\s+OF\s+)?(?:FIGURES|TABLES|ILLUSTRATIONS)|TABLE\s+OF\s+(?:FIGURES|TABLES))\b",
    re.IGNORECASE,
)
_SECTION_SDT_GALLERIES = ("Table of Contents", "Table of Figures", "Bibliographies")
_SHORT_WORDS = 30


def _is_caption(text):
    return bool(text) and bool(_CAPTION_RE.match(text))


def _has_inline_caption(text):
    return bool(text) and bool(_INLINE_CAPTION_RE.search(text))


def _is_heading_paragraph(p):
    style = p.style.name if p.style is not None else ""
    if "Heading" in style:
        return True
    text = p.text.strip()
    if not text:
        return False
    if _HEADING_TEXT_RE.match(text) and len(text) < 60:
        return True
    if len(text) < 60 and text.isupper() and re.search(r"[A-Za-z]", text):
        return True
    return False


def _is_major_heading(p):
    text = p.text.strip()
    if not text or len(text) >= 60:
        return False
    if _MAJOR_HEADING_RE.match(text):
        return True
    if text.isupper() and re.search(r"[A-Za-z]", text):
        return True
    return False


def _is_reference_heading(p):
    text = p.text.strip()
    return bool(text) and len(text) < 60 and bool(_REFERENCE_HEADING_RE.match(text))


def _is_toc_heading(p):
    text = p.text.strip()
    return bool(text) and len(text) < 60 and bool(_TOC_RE.match(text))


def _is_listing_heading(p):
    text = p.text.strip()
    return bool(text) and len(text) < 60 and bool(_LISTING_HEADING_RE.match(text))


def _has_page_break_before(p):
    ppr = p._p.find(qn("w:pPr"))
    if ppr is None:
        return False
    return ppr.find(qn("w:pageBreakBefore")) is not None


def _ends_with_page_break(p):
    runs = p._p.findall(qn("w:r"))
    if not runs:
        return False
    for br in runs[-1].iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _is_short(text):
    return bool(text) and len(text.split()) <= _SHORT_WORDS


def _has_image(p):
    p_elem = p._p
    return (
        p_elem.find(qn("w:drawing")) is not None
        or p_elem.find(qn("w:pict")) is not None
        or any(True for _ in p_elem.iter(qn("pic:pic")))
    )


def _is_table_elem(elem):
    return elem is not None and elem.tag == qn("w:tbl")


def _remove_empty_headings(doc):
    for p in doc.paragraphs:
        style = p.style.name if p.style is not None else ""
        if "Heading" in style and not p.text.strip():
            if _has_image(p):
                continue
            p._p.getparent().remove(p._p)


def _title_page_end(doc):
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        style = p.style.name if p.style is not None else ""
        if style == "Heading 1":
            return i
    for i, p in enumerate(paras):
        if _is_major_heading(p):
            return i
    return None


def apply_pagination(doc):
    _remove_empty_headings(doc)
    title_end = _title_page_end(doc)
    paras = doc.paragraphs
    for idx, p in enumerate(paras):
        pf = p.paragraph_format
        pf.widow_control = True
        text = p.text.strip()

        if (_is_major_heading(p) or _is_reference_heading(p) or _is_toc_heading(p) or _is_listing_heading(p)) and idx > 0:
            j = idx - 1
            while j >= 0 and not paras[j].text.strip():
                j -= 1
            prev = paras[j] if j >= 0 else None
            already_fresh = _has_page_break_before(p) or (
                prev is not None
                and (_has_page_break_before(prev) or _ends_with_page_break(prev))
            )
            if title_end is not None and idx < title_end:
                pass
            else:
                prev_is_heading = prev is not None and _is_heading_paragraph(prev)
                is_boundary = title_end is not None and idx == title_end
                if not already_fresh and (is_boundary or not prev_is_heading):
                    ppr = p._p.get_or_add_pPr()
                    if ppr.find(qn("w:pageBreakBefore")) is None:
                        ppr.insert(0, OxmlElement("w:pageBreakBefore"))

        if _is_heading_paragraph(p) or _is_reference_heading(p) or _is_toc_heading(p) or _is_listing_heading(p):
            pf.keep_with_next = True
            pf.keep_together = True
            continue

        if _is_caption(text):
            pf.keep_with_next = True
            pf.keep_together = True
            continue

        if _is_short(text):
            pf.keep_together = True

        nxt = paras[idx + 1] if idx + 1 < len(paras) else None
        if _has_image(p) and nxt is not None and _is_caption(nxt.text.strip()):
            pf.keep_with_next = True
        elif _is_short(text):
            if nxt is not None and _has_image(nxt):
                pf.keep_with_next = True
            elif _is_table_elem(p._p.getnext()):
                pf.keep_with_next = True

    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))

    _fresh_page_before_section_sdts(doc)
    _remove_blank_page_risks(doc)


def _remove_blank_page_risks(doc):
    paras = doc.paragraphs
    doomed = set()
    for i, p in enumerate(paras):
        ppr = p._p.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None:
            j = i - 1
            while j >= 0 and not paras[j].text.strip():
                doomed.add(id(paras[j]._p))
                j -= 1
    k = len(paras) - 1
    while k >= 0 and not paras[k].text.strip():
        doomed.add(id(paras[k]._p))
        k -= 1
    for p in paras:
        if id(p._p) in doomed:
            p._p.getparent().remove(p._p)


def _sdt_gallery(sdt_elem):
    sdt_pr = sdt_elem.find(qn("w:sdtPr"))
    if sdt_pr is None:
        return None
    obj = sdt_pr.find(qn("w:docPartObj"))
    if obj is None:
        return None
    gal = obj.find(qn("w:docPartGallery"))
    if gal is None:
        return None
    return gal.get(qn("w:val"))


def _fresh_page_before_section_sdts(doc):
    for child in doc.element.body.iterchildren():
        if child.tag != qn("w:sdt"):
            continue
        if _sdt_gallery(child) not in _SECTION_SDT_GALLERIES:
            continue
        content = child.find(qn("w:sdtContent"))
        if content is None:
            continue
        first_p = content.find(qn("w:p"))
        if first_p is None:
            continue
        ppr = first_p.find(qn("w:pPr"))
        if ppr is None:
            ppr = first_p.makeelement(qn("w:pPr"), {})
            first_p.insert(0, ppr)
        if ppr.find(qn("w:pageBreakBefore")) is None:
            ppr.insert(0, OxmlElement("w:pageBreakBefore"))
