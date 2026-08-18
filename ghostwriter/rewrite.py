import io
import logging
import re
from concurrent.futures import ThreadPoolExecutor
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from . import config, llm, citations, voiceprofile

log = logging.getLogger(__name__)
from .pagination import _LISTING_HEADING_RE, _has_inline_caption, _is_caption, _title_page_end, apply_pagination


def is_heading(text):
    stripped = text.strip()
    if not stripped:
        return True
    if stripped.startswith("#") and len(stripped) < 60:
        return True
    if len(stripped) < 60 and re.match(r"^[A-Z][A-Z\s\d\.]+$", stripped):
        return True
    if len(stripped) < 60 and re.match(r"^CHAPTER\s", stripped, re.IGNORECASE):
        return True
    if len(stripped) < 60 and _LISTING_HEADING_RE.match(stripped):
        return True
    if len(stripped) < 60 and re.match(r"^\d+\.\d+", stripped):
        return True
    return False


_REFERENCES_HEADING_RE = re.compile(
    r"^\s*(?:\d+[.)]?\s*)?"
    r"(references?|bibliography|works\s+cited|list\s+of\s+references|"
    r"reference\s+list|citations?)\s*:?\s*$",
    re.IGNORECASE,
)


def is_references_heading(text):
    """True for a heading that opens a reference list."""
    stripped = text.strip().lstrip("#").strip()
    return bool(stripped) and bool(_REFERENCES_HEADING_RE.match(stripped))


def split_document(text):
    paragraphs = []
    for p in re.split(r"\n\n|(?=###\s)", text):
        p = p.strip()
        if len(p) > 10:
            paragraphs.append(p)
    return paragraphs


def extract_heading_prefix(text):
    match = re.match(r"^(###\s+\d+\.\d+\s+[A-Z][A-Z\s]+?)\s+[A-Z][a-z]", text)
    if match:
        return match.group(1).strip()
    return None


KEEP_RATIO = config.KEEP_RATIO
CHUNK_SIZE = config.REWRITE_CHUNK_SIZE

DUP_THRESHOLD = 0.85


def _normalize(s):
    return re.sub(r"[^a-z0-9 ]", " ", s.lower())


def is_duplicate(candidate, seen):
    cand = set(_normalize(candidate).split())
    if len(cand) < 4:
        return False
    for s in seen:
        seen_tokens = set(_normalize(s).split())
        intersection = cand & seen_tokens
        overlap_cand = len(intersection) / max(1, len(cand))
        overlap_seen = len(intersection) / max(1, len(seen_tokens))
        if overlap_cand >= DUP_THRESHOLD or overlap_seen >= DUP_THRESHOLD:
            return True
    return False


def _overlap(a, b):
    ta = set(_normalize(a).split())
    tb = set(_normalize(b).split())
    if not ta:
        return 0.0
    return len(ta & tb) / len(ta)


def _para_is_dup(text, seen):
    cand = set(_normalize(text).split())
    if len(cand) < 10:
        return False
    for s in seen:
        seent = set(_normalize(s).split())
        inter = cand & seent
        if inter and len(inter) / len(cand) >= DUP_THRESHOLD:
            return True
    return False


def split_sentences(text):
    return [s.strip() for s in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\"'`])", text) if s.strip()]


def plan_keep_indices(n_sentences, keep_ratio=KEEP_RATIO):
    if n_sentences < 2 or keep_ratio <= 0:
        return set()
    n_keep = max(1, round(n_sentences * keep_ratio))
    step = n_sentences / n_keep
    return {min(n_sentences - 1, round(k * step)) for k in range(n_keep)}


def _restore_leading_marker(rewritten, original):
    m = re.match(r"^\s*(\d+[.)]?|[a-zA-Z][.)]|\u2022|[-*])\s+", original)
    if not m:
        return rewritten
    marker = m.group(1)
    if rewritten.startswith(marker):
        return rewritten
    return marker + " " + rewritten.strip()


def _parse_numbered(text):
    result = {}
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        m = re.match(r"^\s*(\d+)[.)]?\s+(.*)$", lines[i])
        if m:
            num = int(m.group(1))
            parts = [m.group(2)]
            i += 1
            while i < len(lines):
                if re.match(r"^\s*\d+[.)]?\s+", lines[i]):
                    break
                parts.append(lines[i])
                i += 1
            result[num] = " ".join(parts).strip()
        else:
            i += 1
    return result


MAX_GROWTH = 1.3


def cap_paragraph(text, source_text):
    limit = int(len(source_text.split()) * MAX_GROWTH)
    if len(text.split()) <= limit:
        return text
    source_sentences = set(split_sentences(source_text))
    capped = []
    total = 0
    stopped = False
    for s in split_sentences(text):
        n = len(s.split())
        if s in source_sentences:
            capped.append(s)
            total += n
            continue
        if stopped:
            continue
        if capped and total + n > limit:
            stopped = True
            continue
        capped.append(s)
        total += n
    if not capped:
        return text
    return " ".join(capped)


def cap_document(input_doc, output_text):
    in_paras = input_doc.split("\n\n")
    out_paras = output_text.split("\n\n")
    capped = []
    for i, op in enumerate(out_paras):
        if i < len(in_paras):
            capped.append(cap_paragraph(op, in_paras[i]))
        else:
            capped.append(op)
    return "\n\n".join(normalize_spacing(c) for c in capped)


_TAIL_PATTERNS = [
    re.compile(r",\s+and this is\s+.*?\.?$", re.IGNORECASE),
    re.compile(r",\s+and it is\s+.*?\.?$", re.IGNORECASE),
    re.compile(r",\s+thus\s+[a-z]+ing\s+.*?\.?$", re.IGNORECASE),
    re.compile(r",\s+which is\s+(?:very|an|a|also|integral)\s+.*?\.?$", re.IGNORECASE),
    re.compile(
        r",\s+and\s+.{3,90}?\b(?:inadequate|in need of|capable|very much|"
        r"crucial|important|significant|essential|vital|beneficial|worth|"
        r"necessary|limitation|challenge|problem|often|of such|effective|"
        r"useful|valuable|concerning|problematic|promising)\b.*?\.?$",
        re.IGNORECASE,
    ),
]

_FOOTER_PATTERNS = [
    re.compile(r"^\s*Page\s+\d+\s+of\s+\d+\s+.*$", re.IGNORECASE),
    re.compile(r"^\s*Submission\s+ID\s+trn:oid\s*.*$", re.IGNORECASE),
    re.compile(r"^\s*AI\s+Writing\s+Submission.*$", re.IGNORECASE),
]


def sanitize_footers(text):
    out = []
    for line in text.splitlines():
        if any(p.match(line) for p in _FOOTER_PATTERNS):
            continue
        out.append(line)
    return "\n".join(out)


def normalize_spacing(text):
    return re.sub(r"(\S)\s+([,;:)])", r"\1\2", text)


def strip_added_tail(rewritten, original):
    orig_lower = original.lower()
    cut = None
    for pattern in _TAIL_PATTERNS:
        m = pattern.search(rewritten)
        if m:
            tail = m.group(0)
            if tail.lower() in orig_lower:
                continue
            if cut is None or m.start() < cut:
                cut = m.start()
    if cut is None:
        return rewritten
    candidate = rewritten[:cut].rstrip(" ,")
    if len(candidate.split()) < 3:
        return rewritten
    return candidate


def _rewrite_chunk(chunk, system_prompt):
    lines = []
    for n, s, prev_s, next_s in chunk:
        if prev_s and next_s:
            ctx = '\n   (Between: "%s" ... "%s")' % (prev_s, next_s)
        elif prev_s:
            ctx = '\n   (Previous sentence: "%s")' % prev_s
        elif next_s:
            ctx = '\n   (Next sentence: "%s")' % next_s
        else:
            ctx = ""
        lines.append("%d. %s%s" % (n, s, ctx))
    numbered = "\n".join(lines)
    prompt = f"""Rewrite each numbered sentence below in the author's voice as described in the system prompt.

{numbered}

Rules:
- Produce exactly ONE rewritten sentence per number, in the same order
- Keep the same numbering
- Do not add, remove, merge, or split any sentences
- Say the same idea in clearly DIFFERENT words and sentence structure. This must NOT be a near-copy of the original — change at least half the words and reorder the clauses
- Do NOT add new facts, examples, or ideas
- Do NOT append any new clause that adds evaluation, commentary, or a conclusion (no "and this is...", "thus ...ing", "which is...", or similar)
- Keep all facts, numbers, and any citation exactly as written inside its parentheses, e.g. "(Ismail et al., 2023)"
- Keep any {{CIT_n}} placeholder exactly as written, wherever it appears in the sentence
- Keep any list MARKER exactly ("1", "2.", "•", "-") — it is document structure
- But a "Label: explanation" opening is NOT structure, it is wording. Rewrite it into a flowing sentence in the author's voice rather than preserving the label and colon. "Touchless operation: removes the need to touch surfaces" should become something like "Touchless operation removes the need of touching common surfaces". Never leave a capitalised label followed by a colon
- Carry over every point the original makes — do not summarise, drop a clause, or leave a fact out. But express those points in clearly DIFFERENT words and a different clause order. Completeness of content and freshness of wording are both required; do not trade one for the other by staying close to the original's phrasing
- Use the "(Between: ...)" context only to fit the sentence naturally; never copy words from it into your rewritten sentence
- Self-check before answering: compare each sentence you wrote against the author's example passages in the system prompt. If it sounds too clean, smooth, modern, or AI-like, rewrite it again internally. Also confirm it is NOT a near-copy of the original sentence
- Return only the numbered rewritten sentences, nothing else"""
    result = llm.ask(prompt, system_prompt=system_prompt, temperature=config.REWRITE_TEMPERATURE)
    return _parse_numbered(result)


_profile_cache = None
_profile_loaded = False


def _voice_profile():
    """The measured voice profile, or None if the gate should be skipped.

    Loaded once per process. A missing profile is not an error: the pipeline
    behaves exactly as it did before the gate existed. Regenerate the file with
    ``python -m ghostwriter.voiceprofile <corpus> -o data/style_prompt.txt``.
    """
    global _profile_cache, _profile_loaded
    if _profile_loaded:
        return _profile_cache
    _profile_loaded = True
    if not config.VOICE_GATE_ENABLED:
        return None
    path = Path(config.VOICE_PROFILE_FILE)
    if not path.exists():
        log.info("voice gate off: no profile at %s", path)
        return None
    try:
        _profile_cache = voiceprofile.load_json(path)
        log.info("voice gate on: profile from %s", path)
    except Exception as e:
        log.warning("voice gate off: could not load %s (%s)", path, e)
    return _profile_cache


def _voice_repair_chunk(chunk, system_prompt):
    """Re-prompt a batch of sentences, naming each one's specific breach.

    ``chunk`` is a list of ``(n, original, rewritten, problems)``.
    """
    lines = []
    for n, original, rewritten, problems in chunk:
        lines.append(
            '%d. Original: %s\n   Your rewrite: %s\n   Fix: %s'
            % (n, original, rewritten, "; ".join(problems))
        )
    numbered = "\n".join(lines)
    prompt = f"""Each numbered rewrite below breaks one of the author's voice rules. Fix ONLY the listed problems.

{numbered}

Rules:
- Produce exactly ONE corrected sentence per number, in the same order, keeping the numbering
- Fix only what "Fix:" lists — keep the meaning and everything else about the sentence
- Do NOT drift back toward the original's wording; this must still not be a near-copy
- Do NOT add new facts, examples, or evaluative clauses
- Keep all facts, numbers, and any citation exactly as written; keep any {{CIT_n}} placeholder exactly as written
- If the sentence begins with a list marker or label (e.g. "1", "2.", "•", or "Exponential Capacity Growth:"), keep it exactly at the start
- Return only the numbered corrected sentences, nothing else"""
    result = llm.ask(prompt, system_prompt=system_prompt, temperature=config.REWRITE_TEMPERATURE)
    return _parse_numbered(result)


def _apply_voice_gate(rewrites, original_map, system_prompt):
    """Re-prompt sentences that breach the author's measured voice rules.

    Only per-sentence, deterministic breaches are actionable here — banned
    connectives, forbidden punctuation, length and comma inflation. Distributional
    traits (median, spread) would need sentences merged or split, which the
    one-in-one-out reassembly contract forbids; those are reported instead.

    A repair is accepted only when it reduces the breach count AND stays clear of
    the similarity threshold, so this can never undo the earlier similarity gate.
    """
    profile = _voice_profile()
    if profile is None:
        return

    for attempt in range(config.VOICE_GATE_ATTEMPTS):
        flagged = []
        for n, orig in original_map.items():
            candidate = rewrites.get(n)
            if not candidate:
                continue
            problems = voiceprofile.sentence_violations(profile, candidate, orig)
            if problems:
                flagged.append((n, orig, candidate, problems))

        if not flagged:
            if attempt == 0:
                log.info("voice gate: no breaches")
            return

        log.info("voice gate pass %d: repairing %d sentence(s)", attempt + 1, len(flagged))
        before = {n: len(problems) for n, _, _, problems in flagged}

        chunks = [flagged[i:i + CHUNK_SIZE] for i in range(0, len(flagged), CHUNK_SIZE)]
        workers = min(config.REWRITE_MAX_WORKERS, len(chunks)) or 1

        # This pass only refines an already-complete rewrite, so a provider
        # failure must not discard it: one ConnectTimeout during pass 2 threw
        # away 383 seconds of finished work. Keep what we have and stop.
        try:
            with ThreadPoolExecutor(max_workers=workers) as executor:
                results = list(
                    executor.map(lambda c: _voice_repair_chunk(c, system_prompt), chunks)
                )
        except Exception as e:
            log.warning(
                "voice gate pass %d abandoned (%s: %s); keeping the rewrite as it stands",
                attempt + 1, type(e).__name__, str(e)[:160],
            )
            return

        improved = 0
        for parsed in results:
            for n, fixed in parsed.items():
                orig = original_map.get(n)
                if orig is None or n not in before or not fixed.strip():
                    continue
                if len(voiceprofile.sentence_violations(profile, fixed, orig)) >= before[n]:
                    continue
                if _overlap(fixed, orig) >= 0.7:
                    continue
                rewrites[n] = fixed
                improved += 1

        log.info("voice gate pass %d: accepted %d repair(s)", attempt + 1, improved)
        if not improved:
            return


def _log_document_drift(text):
    """Report distributional drift for the finished document. Reporting only."""
    profile = _voice_profile()
    if profile is None:
        return
    try:
        rows = voiceprofile.document_drift(profile, text)
    except Exception as e:
        log.debug("voice drift report unavailable: %s", e)
        return
    drifted = [r for r in rows if not r[3]]
    if drifted:
        log.info(
            "voice drift: %s",
            "; ".join(f"{label} author={a} output={c}" for label, a, c, _ in drifted),
        )
    else:
        log.info("voice drift: none across %d metrics", len(rows))


def _collapse_line_breaks(text):
    paragraphs = []
    for para in text.split("\n\n"):
        paragraphs.append(re.sub(r"\n", " ", para))
    return "\n\n".join(paragraphs)


def rewrite_document(document, system_prompt):
    document = _collapse_line_breaks(sanitize_footers(document))
    paras = [p for p in document.split("\n\n") if p.strip()]

    # Split off any reference list and re-attach it untouched at the end.
    references = []
    for idx, para in enumerate(paras):
        if is_references_heading(para):
            references = paras[idx:]
            paras = paras[:idx]
            log.info("references section held back (%d paragraphs)", len(references))
            break

    deduped = []
    seen_paras = []
    for para in paras:
        stripped = para.strip()
        if is_heading(stripped) or not _para_is_dup(stripped, seen_paras):
            deduped.append(stripped)
            if not is_heading(stripped):
                seen_paras.append(stripped)
    document = "\n\n".join(deduped)

    seen_kept = []
    to_rewrite = []
    plan = []
    counter = 0

    for para in document.split("\n\n"):
        stripped = para.strip()
        if not stripped:
            plan.append([])
            continue
        if is_heading(stripped):
            plan.append([stripped])
            continue
        sentences = split_sentences(stripped)
        if len(sentences) < 2:
            if len(stripped) < 20:
                plan.append(sentences)
                continue
            counter += 1
            plan.append(["REWRITE:%d" % counter])
            to_rewrite.append((counter, stripped, None, None))
            continue
        keep_indices = plan_keep_indices(len(sentences))
        para_plan = []
        for i, s in enumerate(sentences):
            if i in keep_indices and not is_duplicate(s, seen_kept):
                seen_kept.append(s)
                para_plan.append(s)
            else:
                counter += 1
                para_plan.append("REWRITE:%d" % counter)
                prev_s = sentences[i - 1] if i > 0 else None
                next_s = sentences[i + 1] if i + 1 < len(sentences) else None
                to_rewrite.append((counter, s, prev_s, next_s))
        plan.append(para_plan)

    rewrites = {}
    original_map = {n: s for n, s, _, _ in to_rewrite}
    chunks = [to_rewrite[start:start + CHUNK_SIZE] for start in range(0, len(to_rewrite), CHUNK_SIZE)]
    with ThreadPoolExecutor(max_workers=min(config.REWRITE_MAX_WORKERS, len(chunks)) or 1) as executor:
        for parsed in executor.map(lambda c: _rewrite_chunk(c, system_prompt), chunks):
            rewrites.update(parsed)

    for attempt in range(2):
        too_similar = [
            n for n, orig in original_map.items()
            if rewrites.get(n) and _overlap(rewrites[n], orig) >= 0.7
        ]
        if not too_similar:
            break
        lines = "\n".join(
            "%d. %s" % (n, original_map[n]) for n in too_similar
        )
        prompt = f"""Each numbered sentence below was rewritten but the result is still too close to the original — it looks copied. Rewrite it again with clearly different wording and sentence order, keeping the same meaning and any citation.

{lines}

Rules:
- Change at least half the words; reorder the clauses; do not keep the original's phrasing
- Keep all facts and any citation exactly; keep any {{CIT_n}} placeholder exactly as written
- Keep any list MARKER ("1", "2.", "•", "-"), but rewrite a "Label: explanation" opening into a flowing sentence — do not preserve the label and colon
- Carry over every point the original makes, but in different words — do not shorten, and do not stay close to the original's phrasing
- Do not add new ideas or append evaluative clauses
- Return only the numbered rewritten sentences, nothing else"""
        # Refinement of an already-complete rewrite: a provider failure here
        # should cost the retry, not the document.
        try:
            result = llm.ask(
                prompt, system_prompt=system_prompt, temperature=config.REWRITE_TEMPERATURE
            )
        except Exception as e:
            log.warning(
                "similarity retry %d abandoned (%s: %s); keeping the rewrite as it stands",
                attempt + 1, type(e).__name__, str(e)[:160],
            )
            break
        rewrites.update(_parse_numbered(result))

    # Report what the similarity gate could not fix. Previously these fell
    # through silently, so a sentence left as a near-copy of the source was
    # indistinguishable from one that had been rewritten well.
    unresolved = [
        n for n, orig in original_map.items()
        if rewrites.get(n) and _overlap(rewrites[n], orig) >= 0.7
    ]
    missing = [n for n in original_map if not rewrites.get(n)]
    if unresolved:
        log.warning(
            "similarity gate: %d of %d sentence(s) still >=0.70 overlap with the "
            "source after %d retries", len(unresolved), len(original_map), 2
        )
    if missing:
        log.warning(
            "similarity gate: %d sentence(s) were never returned by the model and "
            "fall back to the ORIGINAL text - lower GHOSTWRITER_CHUNK_SIZE",
            len(missing),
        )

    # The similarity gate above only measures distance from the source sentence.
    # The voice gate measures conformance to the author's own measured habits.
    _apply_voice_gate(rewrites, original_map, system_prompt)

    rebuilt = []
    for para_plan in plan:
        parts = []
        for item in para_plan:
            if isinstance(item, str) and item.startswith("REWRITE:"):
                n = int(item.split(":", 1)[1])
                rewritten = rewrites.get(n, original_map.get(n, ""))
                rewritten = _restore_leading_marker(rewritten, original_map.get(n, ""))
                parts.append(strip_added_tail(rewritten, original_map.get(n, "")))
            else:
                parts.append(item)
        rebuilt.append(" ".join(p for p in parts if p))

    final = cap_document(document, "\n\n".join(rebuilt))
    if references:
        final = final + "\n\n" + "\n\n".join(references)
    _log_document_drift(final)
    return final


def _cell_paragraphs(cell, doc, seen):
    """Paragraphs in a cell, descending into nested tables.

    Merged cells appear once per spanned column in ``row.cells``, so the same
    paragraph would otherwise be yielded several times and rewritten twice.
    """
    for p in cell.paragraphs:
        key = id(p._p)
        if key not in seen:
            seen.add(key)
            yield p
    for nested in cell.tables:
        yield from _table_paragraphs(nested, doc, seen)


def _table_paragraphs(table, doc, seen):
    for row in table.rows:
        for cell in row.cells:
            yield from _cell_paragraphs(cell, doc, seen)


def iter_paragraphs(doc):
    """Every paragraph in document order, including inside table cells.

    ``doc.paragraphs`` returns body-level paragraphs only: anything inside a
    ``<w:tbl>`` is invisible to it. Iterating it alone meant table prose was
    never sent to the model and passed through byte-identical, however it read.
    """
    seen = set()
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            key = id(child)
            if key not in seen:
                seen.add(key)
                yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield from _table_paragraphs(Table(child, doc), doc, seen)


# A table cell is usually a label or a datum, not prose. Paraphrasing those
# damages the table: headers stop matching their columns, "Ultra-High Frequency
# (UHF)" becomes "Ultra-High Frequency which is called UHF.", and terse matrix
# entries turn into sentences. Only rewrite a cell that is genuinely prose.
TABLE_PROSE_MIN_WORDS = 12
TABLE_PROSE_MIN_WORDS_WITH_STOP = 8


def _is_table_prose(text):
    words = text.split()
    if len(words) >= TABLE_PROSE_MIN_WORDS:
        return True
    # A shorter cell still counts if it is punctuated as a sentence.
    return (
        len(words) >= TABLE_PROSE_MIN_WORDS_WITH_STOP
        and re.search(r"[.!?]\s*$", text.strip()) is not None
    )


def _table_header_paragraphs(doc):
    """Paragraph ids in the first row of every table, nested ones included.

    Header cells name the columns. Rewriting them silently breaks the table's
    meaning - "Energy Savings Potential" became "Possible Energy Reduction".
    """
    ids = set()

    def walk(tables):
        for t in tables:
            rows = t.rows
            # A single-row table has no header/data distinction - it is usually a
            # callout or text box, so treating its only row as a header would
            # exempt the whole thing from rewriting.
            if len(rows) >= 2:
                for c in rows[0].cells:
                    for p in c.paragraphs:
                        ids.add(id(p._p))
            for r in rows:
                for c in r.cells:
                    walk(c.tables)

    walk(doc.tables)
    return ids


def _in_table(paragraph):
    parent = paragraph._p.getparent()
    while parent is not None:
        if parent.tag == qn("w:tbl"):
            return True
        parent = parent.getparent()
    return False


def _has_picture(elem):
    return (
        elem.find(qn("w:drawing")) is not None
        or elem.find(qn("w:pict")) is not None
        or any(True for _ in elem.iter(qn("pic:pic")))
    )


def _is_heading_paragraph(p):
    style = p.style.name if p.style is not None else ""
    if "Heading" in style or "Caption" in style:
        return True
    return is_heading(p.text)


def _sentence_bold_info(text, run_spans):
    sents = split_sentences(text)
    if not sents:
        return []
    spans = []
    pos = 0
    for s in sents:
        start = text.find(s, pos)
        if start < 0:
            start = pos
        spans.append((start, start + len(s)))
        pos = start + len(s)
    info = []
    for k, (s0, s1) in enumerate(spans):
        bold_ranges = []
        for r0, r1, b in run_spans:
            if b and r1 > r0:
                a, c = max(s0, r0), min(s1, r1)
                if a < c:
                    bold_ranges.append((a - s0, c - s0))
        info.append((sents[k], bold_ranges))
    return info


def _ranges_cover(ranges, length):
    merged = []
    for a, b in sorted(ranges):
        if merged and a <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], b))
        else:
            merged.append((a, b))
    return bool(merged) and merged[0][0] <= 0 and merged[-1][1] >= length


def _segments_for_partial(out_s, src_s, bold_ranges):
    pos = 0
    pieces = []
    for s0, s1 in bold_ranges:
        bold_txt = src_s[s0:s1]
        idx = out_s.find(bold_txt, pos)
        if idx < 0:
            return None
        if idx > pos:
            pieces.append((out_s[pos:idx], False))
        pieces.append((out_s[idx:idx + len(bold_txt)], True))
        pos = idx + len(bold_txt)
    if pos < len(out_s):
        pieces.append((out_s[pos:], False))
    return pieces


def _replace_paragraph_text(paragraph, new_text, bold_segments=None, citation_records=None):
    p_elem = paragraph._p
    template_rpr = None
    image_children = []
    for child in list(p_elem):
        if child.tag in (
            qn("w:pPr"),
            qn("w:bookmarkStart"),
            qn("w:bookmarkEnd"),
            qn("w:proofErr"),
        ):
            continue
        if _has_picture(child):
            image_children.append(child)
            continue
        if template_rpr is None:
            if child.tag == qn("w:r"):
                rpr = child.find(qn("w:rPr"))
                if rpr is not None:
                    template_rpr = rpr
            elif child.tag == qn("w:hyperlink"):
                for r in child.findall(qn("w:r")):
                    rpr = r.find(qn("w:rPr"))
                    if rpr is not None:
                        template_rpr = rpr
                        break
        p_elem.remove(child)

    if bold_segments is None:
        bold_segments = [(new_text, None)]

    citation_map = {ph: elem for ph, elem, _ in (citation_records or [])}
    used = set()

    new_nodes = []
    for seg_text, bold in bold_segments:
        for piece, ph in citations.split_placeholders(seg_text):
            if ph is not None:
                elem = citation_map.get(ph)
                if elem is not None:
                    new_nodes.append(deepcopy(elem))
                    used.add(ph)
                continue
            if not piece:
                continue
            run = paragraph.add_run(piece)
            rpr = deepcopy(template_rpr) if template_rpr is not None else None
            if bold is not None:
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                b = rpr.find(qn("w:b"))
                if b is None:
                    b = OxmlElement("w:b")
                    rpr.insert(0, b)
                b.set(qn("w:val"), "1" if bold else "0")
            if rpr is not None:
                run._element.insert(0, rpr)
            new_nodes.append(run._element)

    for ph, elem, _ in citation_records or []:
        if ph not in used:
            new_nodes.append(deepcopy(elem))

    for n in new_nodes:
        if n.tag == qn("w:r"):
            p_elem.remove(n)

    anchor = p_elem.find(qn("w:pPr"))
    for bm in p_elem.findall(qn("w:bookmarkStart")):
        anchor = bm
    if anchor is None and len(p_elem):
        anchor = p_elem[-1]
    for n in new_nodes:
        if anchor is not None:
            anchor.addnext(n)
            anchor = n
        else:
            p_elem.insert(0, n)
            anchor = n
    for img in image_children:
        p_elem.append(img)
    return paragraph


_TOC_ENTRY_STYLES = {
    "TableofFigures", "TableofTables", "TOCHeading",
    "TOC1", "TOC2", "TOC3", "TOC4", "TOC5",
    "TOC6", "TOC7", "TOC8", "TOC9", "TOC10",
}


def _is_toc_entry_paragraph(p, fld_tag, instr_tag):
    if next(p.iter(fld_tag), None) is not None:
        return True
    if next(p.iter(instr_tag), None) is not None:
        return True
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        style = pPr.find(qn("w:pStyle"))
        if style is not None and style.get(qn("w:val")) in _TOC_ENTRY_STYLES:
            return True
    text = "".join(t.text or "" for t in p.iter(qn("w:t")))
    return bool(re.match(r"^(Figure|Table|Illustration|Fig\.?)\s*\d+", text.strip()))


def close_unclosed_fields(doc) -> None:
    body = doc.element.body
    para_tag = qn("w:p")
    fld_tag = qn("w:fldChar")
    instr_tag = qn("w:instrText")

    paras = [el for el in body.iter() if el.tag == para_tag]
    para_index = {id(p): i for i, p in enumerate(paras)}

    def _para_of(el):
        parent = el.getparent()
        while parent is not None and parent.tag != para_tag:
            parent = parent.getparent()
        return parent

    open_fields = []
    for el in body.iter():
        if el.tag == fld_tag:
            ftype = el.get(qn("w:fldCharType"))
            if ftype == "begin":
                open_fields.append([_para_of(el), None, None])
            elif ftype == "separate":
                if open_fields:
                    open_fields[-1][1] = _para_of(el)
            elif ftype == "end":
                if open_fields:
                    open_fields.pop()
        elif el.tag == instr_tag and open_fields:
            if open_fields[-1][2] is None:
                text = "".join(el.itertext()).strip()
                if text:
                    open_fields[-1][2] = text

    for begin_para, sep_para, instr in open_fields:
        if sep_para is None:
            sep_para = begin_para
        if sep_para is None:
            continue
        start = para_index.get(id(sep_para), len(paras))
        end_para = sep_para
        for p in paras[start + 1:]:
            if _is_toc_entry_paragraph(p, fld_tag, instr_tag):
                end_para = p
            else:
                break
        r = OxmlElement("w:r")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "end")
        r.append(fld)
        end_para.append(r)


def rewrite_docx(file_bytes: bytes, system_prompt: str) -> bytes:
    doc = Document(io.BytesIO(file_bytes))
    original = list(iter_paragraphs(doc))

    # Everything from a "References"/"Bibliography" heading onward is left
    # exactly as the author wrote it. Reference entries are not prose: journal
    # and conference names cannot be reworded (they came back byte-identical,
    # which reads as "not rewritten"), and when the model did reword them it
    # corrupted real citations - retitling published papers and injecting the
    # author's "in year" coinage into them.
    refs_start = len(original)
    for i, p in enumerate(original):
        if is_references_heading(p.text):
            refs_start = i
            log.info("references section starts at paragraph %d; left verbatim", i)
            break

    header_ids = _table_header_paragraphs(doc)

    raw = []
    cite_counter = [0]
    for i, p in enumerate(original):
        if i >= refs_start:
            break
        visible_text, records, run_spans = citations.extract(p, cite_counter)
        text = re.sub(r"\n", " ", visible_text).strip()
        if not text:
            continue
        if _is_caption(text):
            continue
        if _has_inline_caption(text) or re.search(r"SEQ (?:Figure|Table)", p._p.xml):
            continue
        if _is_heading_paragraph(p):
            continue
        if _in_table(p) and (id(p._p) in header_ids or not _is_table_prose(text)):
            continue
        info = _sentence_bold_info(re.sub(r"\n", " ", visible_text), run_spans)
        raw.append((text, i, info, records))

    merged = []
    consumed = set()
    title_end = _title_page_end(doc) or 0
    for text, i, info, records in raw:
        # A sentence continued across two body paragraphs is worth rejoining;
        # two table cells that merely start lowercase are not the same sentence,
        # so never merge when either side sits in a table.
        mergeable = not _in_table(original[i]) and (
            not merged or not _in_table(original[merged[-1][1]])
        )
        if (
            merged
            and mergeable
            and i >= title_end
            and merged[-1][1] >= title_end
            and not re.search(r"[.!?]\s*$", merged[-1][0])
            and re.match(r"^[a-z]", text)
        ):
            prev_text, prev_i, prev_info, prev_records = merged[-1]
            merged[-1] = (
                prev_text + " " + text,
                prev_i,
                prev_info + info,
                prev_records + records,
            )
            consumed.add(i)
        else:
            merged.append((text, i, info, records))

    body = [(t, i, f, r) for t, i, f, r in merged if len(t) >= 20]

    seen = []
    filtered = []
    for text, i, info, records in body:
        if not _para_is_dup(text, seen):
            filtered.append((text, i, info, records))
            seen.append(text)
    body_paras = [t for t, _, _, _ in filtered]
    body_indices = {i for _, i, _, _ in filtered}
    body_sources = {i: t for t, i, _, _ in filtered}
    body_info = {i: f for _, i, f, _ in filtered}
    body_citations = {i: r for _, i, _, r in filtered}

    rewritten_body = rewrite_document("\n\n".join(body_paras), system_prompt)
    rewritten_paras = [p.strip() for p in rewritten_body.split("\n\n") if p.strip()]

    idx = 0
    for i, p in enumerate(original):
        if i not in body_indices:
            continue
        text = re.sub(r"\n", " ", p.text).strip()
        if not text:
            continue
        rewritten = rewritten_paras[idx] if idx < len(rewritten_paras) else text
        final_text = cap_paragraph(rewritten, body_sources.get(i, text))
        info = body_info.get(i, [])
        records = body_citations.get(i, [])
        out_sents = split_sentences(final_text)
        if info and len(out_sents) <= len(info):
            segments = []
            for n, out_s in enumerate(out_sents):
                src_s, bold_ranges = info[n]
                if not bold_ranges:
                    segments.append((out_s, False))
                elif _ranges_cover(bold_ranges, len(src_s)):
                    segments.append((out_s, True))
                else:
                    pieces = _segments_for_partial(out_s, src_s, bold_ranges)
                    segments.extend(pieces if pieces is not None else [(out_s, True)])
                if n < len(out_sents) - 1:
                    seg_text, seg_bold = segments[-1]
                    segments[-1] = (seg_text + " ", seg_bold)
            _replace_paragraph_text(p, final_text, bold_segments=segments, citation_records=records)
        else:
            _replace_paragraph_text(p, final_text, citation_records=records)
        idx += 1

    for i in sorted(consumed, reverse=True):
        p_elem = original[i]._p
        for child in list(p_elem):
            if child.tag != qn("w:pPr"):
                p_elem.remove(child)

    apply_pagination(doc)
    close_unclosed_fields(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
