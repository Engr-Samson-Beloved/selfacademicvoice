import io
import re
from pathlib import Path

import docx
from pypdf import PdfReader

from .pagination import apply_pagination


def parse_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    return "\n\n".join(paragraphs)


def parse_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def docx_stats(file_bytes: bytes) -> tuple[int, int]:
    doc = docx.Document(io.BytesIO(file_bytes))
    xml = doc.element.body.xml
    images = xml.count("pic:pic")
    tables = len(doc.tables)
    return images, tables


def extract_text(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return parse_docx(file_bytes)
    elif suffix == ".pdf":
        return parse_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .docx or .pdf")


def _make_docx_from_text(text: str) -> bytes:
    doc = docx.Document()
    for para in text.split("\n\n"):
        para = re.sub(r"\n", " ", para).strip()
        if para:
            doc.add_paragraph(para)
    apply_pagination(doc)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
